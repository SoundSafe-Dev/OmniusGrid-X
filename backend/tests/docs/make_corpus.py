#!/usr/bin/env python3
"""
Corpus generator — render a structured SOP model to all five evaluation formats
(md / txt / csv / docx / pdf) from ONE source of truth, so the renderings can
never silently drift apart (the whole point of the eval corpus).

This is a *maintainer* tool. The generated files are committed, so running the
test suite needs nothing from here. Regenerate only when a source model below
changes:

    python3.11 -m venv .venv-rageval
    .venv-rageval/bin/pip install reportlab python-docx   # already backend deps
    .venv-rageval/bin/python backend/tests/docs/make_corpus.py

By design this generator only emits the *newer* documents (e.g. SOP-WH-021).
The original SOP-QA-014 files were hand-authored before this tool existed and
carry carefully-calibrated wording that the query ground truth depends on, so
they are intentionally NOT regenerated here.

A document is an ordered list of typed elements. Each renderer walks the same
list, so md/txt/csv/docx/pdf all carry identical facts. CSV additionally uses
the per-element ``item_id`` / ``param`` / ``role`` / ``ref`` fields to fill the
fixed schema the pipeline's CSV parser expects.
"""

from __future__ import annotations

import csv
import io
import sys
import textwrap
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

HERE = Path(__file__).resolve().parent


# --------------------------------------------------------------------------- #
# Document model — element constructors
# --------------------------------------------------------------------------- #
def meta(pairs: List[Tuple[str, str, str]]) -> Dict[str, Any]:
    # pairs: (item_id, field_label, value)
    return {"kind": "meta", "pairs": pairs}


def section(num: str, title: str) -> Dict[str, Any]:
    return {"kind": "section", "num": num, "title": title}


def subsection(num: str, title: str) -> Dict[str, Any]:
    return {"kind": "subsection", "num": num, "title": title}


def para(text: str, *, item_id: str = "", title: str = "",
         param: str = "", ref: str = "") -> Dict[str, Any]:
    return {"kind": "para", "text": text, "item_id": item_id, "title": title,
            "param": param, "ref": ref}


def defs(items: List[Tuple[str, str, str]]) -> Dict[str, Any]:
    # items: (item_id, term, definition)
    return {"kind": "defs", "items": items}


def responsibilities(items: List[Tuple[str, str, str]]) -> Dict[str, Any]:
    # items: (item_id, role, text)
    return {"kind": "responsibilities", "items": items}


def materials(items: List[Tuple[str, str, str, str]]) -> Dict[str, Any]:
    # items: (item_id, name, detail, param)
    return {"kind": "materials", "items": items}


def steps(sec: str, items: List[Tuple[str, str, str, str]]) -> Dict[str, Any]:
    # items: (n, text, role, param)
    return {"kind": "steps", "section": sec, "items": items}


def corrective(intro: str, items: List[Tuple[str, str, str, str, str]]) -> Dict[str, Any]:
    # items: (item_id, short_title, text, param, ref)
    return {"kind": "corrective", "intro": intro, "items": items}


def record(item_id: str, title: str, text: str, param: str) -> Dict[str, Any]:
    return {"kind": "record", "item_id": item_id, "title": title,
            "text": text, "param": param}


def training(text: str, param: str = "") -> Dict[str, Any]:
    return {"kind": "training", "text": text, "param": param}


def related(items: List[Tuple[str, str]]) -> Dict[str, Any]:
    # items: (code, name)
    return {"kind": "related", "items": items}


def revisions(items: List[Tuple[str, str, str, str]]) -> Dict[str, Any]:
    # items: (rev, date, author, desc)
    return {"kind": "revisions", "items": items}


class Document:
    def __init__(self, *, doc_id: str, basename: str, title: str,
                 subtitle: str, elements: List[Dict[str, Any]]):
        self.doc_id = doc_id
        self.basename = basename
        self.title = title
        self.subtitle = subtitle
        self.elements = elements


# --------------------------------------------------------------------------- #
# The new document: SOP-WH-021 (a different domain than SOP-QA-014, with the
# same structural richness so it exercises the same parser/chunker paths and
# supports the same trap archetypes — near-dup numbers, table-row integrity,
# definition-vs-usage, revision diff, out-of-corpus negatives, scope boundary).
#
# Deliberately kept clear of anything that would answer SOP-QA-014's negative
# queries (no preventive-maintenance schedule; no Line 1/2 cleaning content),
# so both docs can be indexed together without cross-contaminating each other's
# ground truth. Its distinct retention period (12 months vs QA-014's 3 years)
# keeps that fact unambiguous across the corpus.
# --------------------------------------------------------------------------- #
WH_021 = Document(
    doc_id="sop-wh-021",
    basename="SOP-WH-021_Cold_Chain_Temperature_Control",
    title="Refrigerated and Frozen Storage Temperature Control and Cold-Chain Verification",
    subtitle="Distribution Center DC-2 — Cold Rooms A/B and Freezer F1",
    elements=[
        meta([
            ("SOP_Number", "SOP Number", "SOP-WH-021"),
            ("Revision", "Revision", "2.2"),
            ("Effective_Date", "Effective Date", "April 15, 2026"),
            ("Review_Cycle", "Review Cycle", "Annual"),
            ("Document_Owner", "Document Owner", "Warehouse Operations Dept."),
            ("Approved_By", "Approved By", "Distribution Quality Lead"),
            ("Applicable_Areas", "Applicable Area(s)", "DC-2 Cold Rooms A/B, Freezer F1"),
            ("Classification", "Classification", "Internal - Controlled"),
        ]),

        section("1", "Purpose"),
        para(
            "This procedure defines the requirements for monitoring, recording, and "
            "verifying the temperature of refrigerated and frozen product during "
            "receiving, storage, and shipping at Distribution Center DC-2. It is "
            "intended to maintain an unbroken cold chain and to provide a documented, "
            "repeatable basis for shipment hold and release decisions.",
            item_id="Purpose", title="Purpose"),

        section("2", "Scope"),
        para(
            "This SOP applies to all refrigerated and frozen finished goods received, "
            "stored, and shipped through DC-2 Cold Rooms A and B and Freezer F1. It "
            "covers receiving temperature checks, continuous storage monitoring, "
            "temperature excursions, and pre-shipment verification. It does not cover "
            "ambient dry-goods storage, which is addressed in SOP-WH-004.",
            item_id="Scope", title="Scope", ref="SOP-WH-004"),

        section("3", "Definitions"),
        defs([
            ("Cold_Chain", "Cold Chain",
             "The uninterrupted series of temperature-controlled storage and handling "
             "steps that keep a product within its required temperature range from "
             "receipt through shipment."),
            ("TOR", "TOR (Time-Out-of-Refrigeration)",
             "The cumulative time a refrigerated product spends above 40 degrees F "
             "(4 degrees C) while outside controlled cold storage."),
            ("Temperature_Excursion", "Temperature Excursion",
             "Any recorded event in which a storage area's air temperature moves "
             "outside its validated setpoint range for more than 15 minutes."),
            ("Setpoint", "Setpoint",
             "The target air temperature a storage area is controlled to, distinct "
             "from the alarm limit at which corrective action begins."),
            ("Cold_Chain_Release", "Cold-Chain Release",
             "The documented sign-off confirming a shipment met temperature "
             "requirements throughout storage and is cleared to ship."),
            ("Data_Logger", "Data Logger",
             "A calibrated electronic device that records air or product temperature "
             "at fixed intervals."),
        ]),

        section("4", "Responsibilities"),
        responsibilities([
            ("Receiving_Clerk", "Receiving Clerk",
             "Measures and records product temperature at receipt and completes the "
             "Receiving Temperature Log before product is moved to storage."),
            ("Cold_Storage_Operator", "Cold Storage Operator",
             "Monitors area temperatures, responds to alarms, and records corrective "
             "actions on the Excursion Log."),
            ("Warehouse_QA_Technician", "Warehouse QA Technician",
             "Performs pre-shipment temperature verification and approves or rejects "
             "shipments for cold-chain release."),
            ("Shift_Lead", "Shift Lead",
             "Confirms the shipping schedule, resolves cold-dock conflicts, and "
             "co-signs the Cold-Chain Release Record."),
            ("Distribution_Quality_Lead", "Distribution Quality Lead",
             "Owns this SOP, reviews excursion and deviation reports, and approves any "
             "temporary variance to stated parameters."),
        ]),

        section("5", "Materials and Equipment"),
        materials([
            ("Data_Loggers", "Calibrated data loggers",
             "One per storage area, logging air temperature at 15-minute intervals",
             "15-minute interval"),
            ("Probe_Thermometer", "Handheld probe thermometer",
             "Calibrated; used for product core temperature checks", ""),
            ("Cold_Rooms", "Cold Rooms A and B",
             "Validated setpoint 36 degrees F (2 degrees C), acceptable range 34-40 degrees F",
             "34-40 degrees F"),
            ("Freezer_F1", "Freezer F1",
             "Validated setpoint -10 degrees F, acceptable range -20 to 0 degrees F",
             "-20 to 0 degrees F"),
            ("Forms", "Cold-chain forms",
             "Excursion Log, Receiving Temperature Log, and Cold-Chain Release Record", ""),
        ]),

        section("6", "Procedure"),
        subsection("6.1", "Receiving Inspection"),
        steps("6.1", [
            ("1", "Receiving Clerk verifies that the carrier's transport temperature "
                  "record accompanies the shipment.", "Receiving Clerk", ""),
            ("2", "Measure the product core temperature of at least three cases per "
                  "pallet using the calibrated probe thermometer.", "Receiving Clerk",
             "min 3 cases/pallet"),
            ("3", "Refrigerated product must be at or below 40 degrees F to be accepted "
                  "into Cold Rooms A/B; frozen product must be at or below 0 degrees F "
                  "to be accepted into Freezer F1.", "Receiving Clerk",
             "refrigerated <=40 F; frozen <=0 F"),
            ("4", "Reject and quarantine any refrigerated shipment with a product core "
                  "temperature above 45 degrees F; place it on hold pending "
                  "disposition.", "Receiving Clerk", "reject > 45 degrees F"),
            ("5", "Record all receiving temperatures on the Receiving Temperature Log "
                  "before product is moved to storage.", "Receiving Clerk", ""),
        ]),

        subsection("6.2", "Storage Monitoring"),
        steps("6.2", [
            ("6", "Each storage area is continuously monitored by a calibrated data "
                  "logger recording air temperature every 15 minutes.",
             "Cold Storage Operator", "air temp every 15 minutes"),
            ("7", "Cold Storage Operator reviews logger readings at the start of every "
                  "shift and records the review on the Storage Monitoring Log.",
             "Cold Storage Operator", ""),
            ("8", "Product core temperature is spot-checked and recorded every 4 hours "
                  "for product staged on the cold dock.", "Cold Storage Operator",
             "core temp every 4 hours"),
            ("9", "The cumulative time-out-of-refrigeration (TOR) for any refrigerated "
                  "product staged outside cold storage must not exceed 2 hours.",
             "Cold Storage Operator", "TOR <= 2 hours"),
        ]),

        subsection("6.3", "Temperature Excursions"),
        steps("6.3", [
            ("10", "An area temperature outside its validated range for more than 15 "
                   "minutes is a Temperature Excursion and triggers an alarm.",
             "Cold Storage Operator", "> 15 minutes out of range"),
            ("11", "Cold Storage Operator responds to the alarm, identifies the cause, "
                   "and records the event and corrective action on the Excursion Log.",
             "Cold Storage Operator", ""),
            ("12", "Product exposed to an excursion is placed on hold and evaluated "
                   "for disposition before it may ship.", "Warehouse QA Technician", ""),
        ]),

        subsection("6.4", "Pre-Shipment Verification"),
        steps("6.4", [
            ("13", "Warehouse QA Technician reviews the storage monitoring record for "
                   "the shipment's storage period and confirms no unresolved excursion "
                   "applies.", "Warehouse QA Technician", ""),
            ("14", "Measure and record the product core temperature of a minimum of "
                   "five cases from the outbound shipment.", "Warehouse QA Technician",
             "minimum 5 cases"),
            ("15", "A refrigerated shipment passes when every measured case is at or "
                   "below 40 degrees F; any case above 40 degrees F requires the "
                   "shipment to be held and re-evaluated before release.",
             "Warehouse QA Technician", "pass at or below 40 degrees F"),
            ("16", "Warehouse QA Technician completes and signs the Cold-Chain Release "
                   "Record; Shift Lead co-signs before the shipment may leave the "
                   "dock.", "Warehouse QA Technician, Shift Lead", ""),
        ]),

        section("7", "Corrective Actions"),
        corrective(
            "If any verification or monitoring step fails, the following actions apply:",
            [
                ("CA1", "Product hold on failure",
                 "Product affected by an excursion or a failed verification is placed "
                 "on hold and evaluated for disposition before release.", "", ""),
                ("CA2", "Two consecutive failures",
                 "Two consecutive failed pre-shipment verifications on the same "
                 "shipment trigger a documented deviation report and notification to "
                 "the Distribution Quality Lead before a third attempt.",
                 "2 consecutive failures", ""),
                ("CA3", "Ship-before-pass recall",
                 "Any shipment loaded before a passing verification is recorded is "
                 "recalled from the dock and placed on hold.", "", ""),
                ("CA4", "Repeated excursions",
                 "Repeated excursions (3 or more in a rolling 30-day period) in the "
                 "same storage area trigger a root-cause investigation per SOP-WH-030.",
                 "3+ events / 30 days", "SOP-WH-030"),
            ]),

        section("8", "Documentation and Record Retention"),
        record("Retention", "Record Retention",
               "The Receiving Temperature Log, Storage Monitoring Log, Excursion Log, "
               "data logger exports, and Cold-Chain Release Records generated by this "
               "procedure are retained by Warehouse Operations for a minimum of 12 "
               "months from the date of creation.",
               "Minimum 12 months"),

        section("9", "Training Requirements"),
        training(
            "All personnel performing tasks under this SOP must complete cold-chain "
            "handling and temperature-monitoring training and demonstrate competency "
            "before working unsupervised. Refresher training is required annually and "
            "whenever this SOP is revised at a level that changes a stated parameter "
            "(temperature, time, or interval).",
            "Annual refresher"),

        section("10", "Related Documents"),
        related([
            ("SOP-WH-004", "Ambient Dry-Goods Storage"),
            ("SOP-WH-030", "Root Cause Investigation, Distribution"),
            ("SOP-SN-005", "Warehouse Pest Control Program"),
            ("SOP-QA-011", "Product Disposition and Hold Procedure"),
            ("Attachment A", "Receiving Temperature Log (form)"),
            ("Attachment B", "Storage Area Setpoint and Alarm Limits"),
        ]),

        section("11", "Revision History"),
        revisions([
            ("2.0", "2024-05-02", "Warehouse Ops",
             "Initial consolidated cold-chain SOP for DC-2, replacing area work "
             "instructions."),
            ("2.1", "2025-07-18", "Warehouse Ops",
             "Added continuous data-logger monitoring at 15-minute intervals and the "
             "Excursion Log."),
            ("2.2", "2026-04-15", "Warehouse Ops",
             "Lowered the refrigerated receiving rejection threshold from 47 to 45 "
             "degrees F."),
        ]),
    ],
)

DOCUMENTS = [WH_021]


# --------------------------------------------------------------------------- #
# Renderers
# --------------------------------------------------------------------------- #
def _section_num(el: Dict[str, Any], current: str) -> str:
    return current


def render_md(doc: Document) -> str:
    out: List[str] = [f"# {doc.title}", f"*{doc.subtitle}*", ""]
    for el in doc.elements:
        k = el["kind"]
        if k == "meta":
            out += ["| Field | Value | Field | Value |", "|---|---|---|---|"]
            p = el["pairs"]
            for i in range(0, len(p), 2):
                left = p[i]
                right = p[i + 1] if i + 1 < len(p) else ("", "", "")
                out.append(f"| **{left[1]}** | {left[2]} | "
                           f"{('**' + right[1] + '**') if right[1] else ''} | {right[2]} |")
            out.append("")
        elif k == "section":
            out += [f"## {el['num']}. {el['title']}", ""]
        elif k == "subsection":
            out += [f"### {el['num']} {el['title']}", ""]
        elif k == "para":
            out += [el["text"], ""]
        elif k == "defs":
            for _id, term, d in el["items"]:
                out.append(f"- **{term}:** {d}")
            out.append("")
        elif k == "responsibilities":
            out += ["| Role | Responsibilities |", "|---|---|"]
            for _id, role, text in el["items"]:
                out.append(f"| {role} | {text} |")
            out.append("")
        elif k == "materials":
            for _id, name, detail, _param in el["items"]:
                out.append(f"- {name}: {detail}")
            out.append("")
        elif k == "steps":
            for n, text, _role, _param in el["items"]:
                out.append(f"{n}. {text}")
            out.append("")
        elif k == "corrective":
            out += [el["intro"], ""]
            for _id, _short, text, _param, _ref in el["items"]:
                out.append(f"- {text}")
            out.append("")
        elif k == "record":
            out += [el["text"], ""]
        elif k == "training":
            out += [el["text"], ""]
        elif k == "related":
            for code, name in el["items"]:
                out.append(f"- {code} – {name}")
            out.append("")
        elif k == "revisions":
            out += ["| Rev | Date | Author | Description of Change |", "|---|---|---|---|"]
            for rev, date, author, desc in el["items"]:
                out.append(f"| {rev} | {date} | {author} | {desc} |")
            out.append("")
    out += ["*— End of Document —*", ""]
    return "\n".join(out)


def _wrap(text: str, width: int = 78, indent: str = "", subsequent: str = "") -> str:
    return textwrap.fill(text, width=width, initial_indent=indent,
                         subsequent_indent=subsequent)


def render_txt(doc: Document) -> str:
    bar = "=" * 80
    dash = "-" * 80
    out: List[str] = [bar, doc.title.upper(), doc.subtitle, bar, ""]
    for el in doc.elements:
        k = el["kind"]
        if k == "meta":
            width = max(len(f[1]) for f in el["pairs"]) + 2
            for _id, field, value in el["pairs"]:
                out.append(f"{(field + ':').ljust(width)} {value}")
            out.append("")
        elif k == "section":
            out += [dash, f"{el['num']}. {el['title'].upper()}", dash]
        elif k == "subsection":
            out += ["", f"{el['num']} {el['title']}", ""]
        elif k == "para":
            out += [_wrap(el["text"]), ""]
        elif k == "defs":
            for _id, term, d in el["items"]:
                out.append(_wrap(f"{term}: {d}", indent="- ", subsequent="  "))
            out.append("")
        elif k == "responsibilities":
            for _id, role, text in el["items"]:
                out.append(_wrap(f"{role}: {text}", indent="- ", subsequent="  "))
            out.append("")
        elif k == "materials":
            for _id, name, detail, _param in el["items"]:
                out.append(_wrap(f"{name}: {detail}", indent="- ", subsequent="  "))
            out.append("")
        elif k == "steps":
            for n, text, _role, _param in el["items"]:
                out.append(_wrap(f"{n}. {text}", subsequent="   "))
            out.append("")
        elif k == "corrective":
            out += [_wrap(el["intro"]), ""]
            for _id, _short, text, _param, _ref in el["items"]:
                out.append(_wrap(text, indent="- ", subsequent="  "))
            out.append("")
        elif k == "record":
            out += [_wrap(el["text"]), ""]
        elif k == "training":
            out += [_wrap(el["text"]), ""]
        elif k == "related":
            for code, name in el["items"]:
                out.append(f"- {code} – {name}")
            out.append("")
        elif k == "revisions":
            for rev, date, author, desc in el["items"]:
                out.append(_wrap(f"Rev {rev} ({date}, {author}): {desc}",
                                 subsequent="  "))
            out.append("")
    out += ["", "-- End of Document --", ""]
    return "\n".join(out)


def render_csv(doc: Document) -> str:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Row_Type", "Section", "Item_ID", "Title", "Content",
                "Parameter_Value", "Responsible_Role", "Reference"])
    cur = "0"
    cur_title = ""
    for el in doc.elements:
        k = el["kind"]
        if k == "meta":
            for _id, field, value in el["pairs"]:
                w.writerow(["Metadata", "0", _id, field, value, "", "", ""])
        elif k == "section":
            cur, cur_title = el["num"], el["title"]
        elif k == "subsection":
            cur = el["num"]
        elif k == "para":
            w.writerow(["Narrative", cur, el.get("item_id") or cur_title,
                        el.get("title") or cur_title, el["text"], el.get("param", ""),
                        "", el.get("ref", "")])
        elif k == "defs":
            for _id, term, d in el["items"]:
                w.writerow(["Definition", "3", _id, term, d, "", "", ""])
        elif k == "responsibilities":
            for _id, role, text in el["items"]:
                w.writerow(["Responsibility", "4", _id, role, text, "", role, ""])
        elif k == "materials":
            for _id, name, detail, param in el["items"]:
                w.writerow(["Material", "5", _id, name, detail, param, "", ""])
        elif k == "steps":
            for n, text, role, param in el["items"]:
                w.writerow(["Procedure_Step", el["section"], n, f"Step {n}",
                            text, param, role, ""])
        elif k == "corrective":
            for _id, short, text, param, ref in el["items"]:
                w.writerow(["Corrective_Action", "7", _id, short, text, param, "", ref])
        elif k == "record":
            w.writerow(["Record", "8", el["item_id"], el["title"], el["text"],
                        el["param"], "", ""])
        elif k == "training":
            w.writerow(["Training", "9", "Training_Req", "Training Requirements",
                        el["text"], el.get("param", ""), "", ""])
        elif k == "related":
            for code, name in el["items"]:
                w.writerow(["Related_Document", "10", code, name, name, "", "", code])
        elif k == "revisions":
            for rev, date, author, desc in el["items"]:
                w.writerow(["Revision", "11", rev, f"Rev {rev} ({date})", desc,
                            date, "", ""])
    return buf.getvalue()


def render_docx(doc: Document, path: Path) -> None:
    import docx
    from docx.shared import Pt

    d = docx.Document()
    d.add_heading(doc.title, level=0)
    d.add_paragraph(doc.subtitle)
    for el in doc.elements:
        k = el["kind"]
        if k == "meta":
            t = d.add_table(rows=0, cols=2)
            t.style = "Table Grid"
            for _id, field, value in el["pairs"]:
                row = t.add_row().cells
                row[0].text = field
                row[1].text = value
        elif k == "section":
            d.add_heading(f"{el['num']}. {el['title']}", level=1)
        elif k == "subsection":
            d.add_heading(f"{el['num']} {el['title']}", level=2)
        elif k == "para":
            d.add_paragraph(el["text"])
        elif k == "defs":
            for _id, term, dd in el["items"]:
                p = d.add_paragraph(style="List Bullet")
                run = p.add_run(f"{term}: ")
                run.bold = True
                p.add_run(dd)
        elif k == "responsibilities":
            t = d.add_table(rows=1, cols=2)
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            hdr[0].text, hdr[1].text = "Role", "Responsibilities"
            for _id, role, text in el["items"]:
                row = t.add_row().cells
                row[0].text, row[1].text = role, text
        elif k == "materials":
            for _id, name, detail, _param in el["items"]:
                d.add_paragraph(f"{name}: {detail}", style="List Bullet")
        elif k == "steps":
            for n, text, _role, _param in el["items"]:
                d.add_paragraph(f"{n}. {text}", style="List Number")
        elif k == "corrective":
            d.add_paragraph(el["intro"])
            for _id, _short, text, _param, _ref in el["items"]:
                d.add_paragraph(text, style="List Bullet")
        elif k == "record":
            d.add_paragraph(el["text"])
        elif k == "training":
            d.add_paragraph(el["text"])
        elif k == "related":
            for code, name in el["items"]:
                d.add_paragraph(f"{code} – {name}", style="List Bullet")
        elif k == "revisions":
            t = d.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            for i, h in enumerate(["Rev", "Date", "Author", "Description of Change"]):
                hdr[i].text = h
            for rev, date, author, desc in el["items"]:
                row = t.add_row().cells
                row[0].text, row[1].text, row[2].text, row[3].text = rev, date, author, desc
    d.save(str(path))


def render_pdf(doc: Document, path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer, Table,
                                    TableStyle)

    ss = getSampleStyleSheet()
    # Font sizes drive the PDF parser's heading detection heuristic, so keep
    # title > h1 > h2 > body clearly separated.
    title_s = ParagraphStyle("Title2", parent=ss["Title"], fontSize=18, spaceAfter=4)
    sub_s = ParagraphStyle("Sub", parent=ss["Normal"], fontSize=11,
                           textColor=colors.grey, spaceAfter=12)
    h1_s = ParagraphStyle("H1", parent=ss["Heading1"], fontSize=15, spaceBefore=12, spaceAfter=6)
    h2_s = ParagraphStyle("H2", parent=ss["Heading2"], fontSize=13, spaceBefore=8, spaceAfter=4)
    body_s = ParagraphStyle("Body", parent=ss["Normal"], fontSize=10.5, leading=14, spaceAfter=6)
    bullet_s = ParagraphStyle("Bul", parent=body_s, leftIndent=14, bulletIndent=4)

    story: List[Any] = [Paragraph(doc.title, title_s), Paragraph(doc.subtitle, sub_s)]
    tbl_style = TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
    ])

    def cell(text: str) -> Paragraph:
        return Paragraph(text, ParagraphStyle("Cell", parent=body_s, fontSize=9.5, spaceAfter=0))

    for el in doc.elements:
        k = el["kind"]
        if k == "meta":
            rows = [[cell(f"<b>{f}</b>"), cell(v)] for _id, f, v in el["pairs"]]
            t = Table(rows, colWidths=[2.1 * inch, 4.4 * inch])
            t.setStyle(tbl_style)
            story += [t, Spacer(1, 10)]
        elif k == "section":
            story.append(Paragraph(f"{el['num']}. {el['title']}", h1_s))
        elif k == "subsection":
            story.append(Paragraph(f"{el['num']} {el['title']}", h2_s))
        elif k == "para":
            story.append(Paragraph(el["text"], body_s))
        elif k == "defs":
            for _id, term, dd in el["items"]:
                story.append(Paragraph(f"<b>{term}:</b> {dd}", bullet_s))
            story.append(Spacer(1, 6))
        elif k == "responsibilities":
            rows = [[cell("<b>Role</b>"), cell("<b>Responsibilities</b>")]]
            rows += [[cell(role), cell(text)] for _id, role, text in el["items"]]
            t = Table(rows, colWidths=[1.9 * inch, 4.6 * inch])
            t.setStyle(tbl_style)
            story += [t, Spacer(1, 10)]
        elif k == "materials":
            for _id, name, detail, _param in el["items"]:
                story.append(Paragraph(f"• {name}: {detail}", bullet_s))
            story.append(Spacer(1, 6))
        elif k == "steps":
            for n, text, _role, _param in el["items"]:
                story.append(Paragraph(f"{n}. {text}", body_s))
        elif k == "corrective":
            story.append(Paragraph(el["intro"], body_s))
            for _id, _short, text, _param, _ref in el["items"]:
                story.append(Paragraph(f"• {text}", bullet_s))
            story.append(Spacer(1, 6))
        elif k == "record":
            story.append(Paragraph(el["text"], body_s))
        elif k == "training":
            story.append(Paragraph(el["text"], body_s))
        elif k == "related":
            for code, name in el["items"]:
                story.append(Paragraph(f"• {code} – {name}", bullet_s))
            story.append(Spacer(1, 6))
        elif k == "revisions":
            rows = [[cell("<b>Rev</b>"), cell("<b>Date</b>"), cell("<b>Author</b>"),
                     cell("<b>Description of Change</b>")]]
            rows += [[cell(rev), cell(date), cell(author), cell(desc)]
                     for rev, date, author, desc in el["items"]]
            t = Table(rows, colWidths=[0.6 * inch, 1.0 * inch, 1.3 * inch, 3.6 * inch])
            t.setStyle(tbl_style)
            story += [t, Spacer(1, 10)]

    SimpleDocTemplate(str(path), pagesize=letter,
                      topMargin=0.7 * inch, bottomMargin=0.7 * inch,
                      leftMargin=0.8 * inch, rightMargin=0.8 * inch).build(story)


# --------------------------------------------------------------------------- #
# Driver
# --------------------------------------------------------------------------- #
def generate(doc: Document, out_dir: Path) -> List[Path]:
    written: List[Path] = []
    md_path = out_dir / f"{doc.basename}.md"
    txt_path = out_dir / f"{doc.basename}.txt"
    csv_path = out_dir / f"{doc.basename}.csv"
    md_path.write_text(render_md(doc), encoding="utf-8")
    txt_path.write_text(render_txt(doc), encoding="utf-8")
    csv_path.write_text(render_csv(doc), encoding="utf-8")
    written += [md_path, txt_path, csv_path]
    render_docx(doc, out_dir / f"{doc.basename}.docx")
    render_pdf(doc, out_dir / f"{doc.basename}.pdf")
    written += [out_dir / f"{doc.basename}.docx", out_dir / f"{doc.basename}.pdf"]
    return written


def main() -> int:
    out_dir = HERE
    for doc in DOCUMENTS:
        paths = generate(doc, out_dir)
        for p in paths:
            print(f"  wrote {p.relative_to(HERE.parent.parent.parent)}  ({p.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
