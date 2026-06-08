"""
DOCX Parser Service

Extracts structure (heading hierarchy, sections, tables, metadata) from Word
.docx files so multi-section documents can be converted into
CorrelationScenarios, mirroring the multi-tab spreadsheet intake pipeline.

Uses python-docx. The output structure is consumed by
``document_domain_mapper`` and ``document_scenario_builder``.
"""

from typing import Dict, List, Any, Optional
import io
import re

import structlog

from app.services.shared_key_detector import (
    extract_keys_from_text,
    extract_keys_from_filename,
    extract_keys_from_metadata,
)

logger = structlog.get_logger()


DEFAULT_MAX_SECTIONS = 100
DEFAULT_MAX_PARAGRAPHS_PER_SECTION = 50


def estimate_processing_seconds(section_count: int) -> float:
    """Estimate ingestion+analysis time. ~0.3s per section."""
    return round(0.3 * max(section_count, 1), 1)


def _is_heading(style_name: Optional[str]) -> Optional[int]:
    """Return heading level (1-9) if the paragraph style is a heading."""
    if not style_name:
        return None
    m = re.match(r"heading\s*(\d)", style_name.strip().lower())
    if m:
        return int(m.group(1))
    if style_name.strip().lower() in ("title",):
        return 0
    return None


def _table_to_rows(table) -> List[List[str]]:
    rows: List[List[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _normalize_meta(core_props) -> Dict[str, Any]:
    meta: Dict[str, Any] = {}
    if not core_props:
        return meta
    for key in ("title", "author", "subject", "category", "comments",
                "created", "modified", "last_modified_by"):
        try:
            value = getattr(core_props, key, None)
            if value:
                meta[key] = str(value)
        except Exception:
            continue
    return meta


def parse_docx_structure(
    content: bytes,
    filename: str,
    max_sections: int = DEFAULT_MAX_SECTIONS,
) -> Dict[str, Any]:
    """
    Parse a .docx into a structured representation.

    Returns:
        {
          "type": "report",
          "subtype": "docx",
          "document_metadata": {...},
          "section_count": int,
          "sections_parsed": int,
          "truncated": bool,
          "sections": [
            {"section_id", "heading", "level", "paragraphs", "tables", "shared_keys"}
          ],
          "tables": [...],          # flattened with section_id
          "shared_keys": [...],
          "estimated_seconds": float,
        }
    """
    try:
        import docx
    except ImportError as e:  # pragma: no cover - dependency guard
        logger.error("python_docx_missing", error=str(e))
        raise RuntimeError("python-docx is required for DOCX parsing") from e

    document = docx.Document(io.BytesIO(content))
    document_metadata = _normalize_meta(getattr(document, "core_properties", None))

    sections: List[Dict[str, Any]] = []
    all_tables: List[Dict[str, Any]] = []

    # Walk paragraphs, grouping into sections by heading boundaries.
    current = {
        "section_id": 0,
        "heading": document_metadata.get("title", "Document Start"),
        "level": 0,
        "paragraphs": [],
        "tables": [],
    }
    truncated = False

    for para in document.paragraphs:
        text = (para.text or "").strip()
        level = _is_heading(getattr(para.style, "name", None))
        if level is not None and text:
            # Close current section and start a new one.
            if current["paragraphs"] or current["tables"]:
                sections.append(current)
                if len(sections) >= max_sections:
                    truncated = True
                    break
            current = {
                "section_id": len(sections),
                "heading": text,
                "level": level,
                "paragraphs": [],
                "tables": [],
            }
        elif text:
            if len(current["paragraphs"]) < DEFAULT_MAX_PARAGRAPHS_PER_SECTION:
                current["paragraphs"].append(text)

    if not truncated and (current["paragraphs"] or current["tables"] or not sections):
        sections.append(current)

    # Attach tables (python-docx exposes tables separately, not inline-ordered).
    for t_idx, table in enumerate(document.tables):
        rows = _table_to_rows(table)
        target_section = sections[-1]["section_id"] if sections else 0
        all_tables.append({"section_id": target_section, "table_index": t_idx, "rows": rows})
        if sections:
            sections[-1]["tables"].append(rows)

    # Per-section shared keys.
    for sec in sections:
        blob = " ".join(sec["paragraphs"])
        for rows in sec["tables"]:
            for row in rows:
                blob += " " + " ".join(str(c) for c in row)
        sec["shared_keys"] = extract_keys_from_text(blob)

    # Aggregate shared keys: filename + metadata + content.
    shared_keys: List[str] = []
    shared_keys.extend(extract_keys_from_filename(filename))
    shared_keys.extend(extract_keys_from_metadata(document_metadata))
    for sec in sections:
        shared_keys.extend(sec["shared_keys"])
    shared_keys = list(dict.fromkeys([k for k in shared_keys if k]))

    return {
        "type": "report",
        "subtype": "docx",
        "document_metadata": document_metadata,
        "section_count": len(sections),
        "sections_parsed": len(sections),
        "truncated": truncated,
        "sections": sections,
        "tables": all_tables,
        "shared_keys": shared_keys,
        "estimated_seconds": estimate_processing_seconds(len(sections)),
    }
